"""
DOCX处理模块
批量填充DOCX邀请函模板并生成高清PNG图片
"""
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import Dict, Any
from utils import find_soffice
import subprocess
import sys


class DOCXHandler:
    """处理DOCX文档"""

    def __init__(self, template_path: str, output_dir: str):
        """初始化DOCX处理器

        Args:
            template_path: DOCX模板文件路径
            output_dir: 输出目录路径
        """
        self.template_path = template_path
        self.output_dir = output_dir

    def fill_template(self, data: Dict[str, Any], output_name: str) -> str:
        """使用数据填充DOCX模板

        Args:
            data: 包含填充数据的字典
            output_name: 输出文件名（不含扩展名）

        Returns:
            生成的DOCX文件路径
        """
        try:
            # 加载模板
            doc = Document(self.template_path)

            def replace_placeholders(text: str) -> str:
                if text is None:
                    return text
                replaced = str(text)

                # 标准字段替换
                replaced = replaced.replace("{{name}}", str(data.get("姓名", "")))
                replaced = replaced.replace("{name}", str(data.get("姓名", "")))
                replaced = replaced.replace("{{unit}}", str(data.get("单位", "")))
                replaced = replaced.replace("{unit}", str(data.get("单位", "")))
                replaced = replaced.replace("{{position}}", str(data.get("职务", "")))
                replaced = replaced.replace("{position}", str(data.get("职务", "")))

                for key, value in data.items():
                    if key is None:
                        continue
                    replaced = replaced.replace(f"{{{{{key}}}}}", str(value or ""))
                    replaced = replaced.replace(f"{{{key}}}", str(value or ""))
                return replaced

            placeholder_pattern = re.compile(r"\{\{[^}]+\}\}|\{[^}]+\}")

            def get_run_segments(paragraph):
                runs = list(paragraph.runs)
                positions = []
                offset = 0
                for run in runs:
                    start = offset
                    end = offset + len(run.text)
                    positions.append((start, end, run))
                    offset = end
                return positions

            W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

            def get_run_r_fonts(run):
                r_pr = run._element.rPr
                if r_pr is None:
                    return None
                return r_pr.find(W_NS + 'rFonts')

            def set_run_font(target_run, ascii_name=None, east_asia_name=None, h_ansi_name=None, cs_name=None):
                r_pr = target_run._element.get_or_add_rPr()
                r_fonts = r_pr.find(W_NS + 'rFonts')
                if r_fonts is None:
                    from lxml import etree
                    r_fonts = etree.SubElement(r_pr, W_NS + 'rFonts')
                if ascii_name is not None:
                    r_fonts.set(qn('w:ascii'), ascii_name)
                if east_asia_name is not None:
                    r_fonts.set(qn('w:eastAsia'), east_asia_name)
                if h_ansi_name is not None:
                    r_fonts.set(qn('w:hAnsi'), h_ansi_name)
                if cs_name is not None:
                    r_fonts.set(qn('w:cs'), cs_name)
                # Also set font.name for compatibility
                if ascii_name is not None:
                    target_run.font.name = ascii_name

            def set_run_underline(target_run, underline_value, source_run=None):
                r_pr = target_run._element.get_or_add_rPr()
                
                # Check if source had an explicit underline element
                if source_run is not None:
                    src_rpr = source_run._element.rPr
                    src_u = src_rpr.find(W_NS + 'u') if src_rpr is not None else None
                    has_explicit_underline = src_u is not None
                else:
                    has_explicit_underline = (underline_value is not None and underline_value is not False)
                
                if not has_explicit_underline:
                    # Source had no <w:u> element — don't add one, let it inherit from paragraph style
                    return
                
                u_elem = r_pr.find(W_NS + 'u')
                if u_elem is None:
                    from lxml import etree
                    u_elem = etree.SubElement(r_pr, W_NS + 'u')
                if underline_value is None or underline_value is False:
                    u_elem.set(qn('w:val'), 'none')
                else:
                    u_elem.set(qn('w:val'), 'single')

            def copy_run_style(source_run, target_run):
                """复制源run的所有样式到目标run，包括eastAsia字体"""
                font = source_run.font
                target_font = target_run.font

                try:
                    r_fonts = get_run_r_fonts(source_run)
                    if r_fonts is not None:
                        set_run_font(
                            target_run,
                            ascii_name=r_fonts.get(qn('w:ascii')),
                            east_asia_name=r_fonts.get(qn('w:eastAsia')),
                            h_ansi_name=r_fonts.get(qn('w:hAnsi')),
                            cs_name=r_fonts.get(qn('w:cs')),
                        )
                    else:
                        # No rFonts in source, use font.name as fallback
                        if font.name is not None:
                            target_font.name = font.name
                    if font.size is not None:
                        target_font.size = font.size
                    if font.color and font.color.rgb is not None:
                        target_font.color.rgb = font.color.rgb
                    if font.highlight_color is not None:
                        target_font.highlight_color = font.highlight_color
                except Exception:
                    pass

                try:
                    target_run.bold = font.bold
                    target_run.italic = font.italic
                    set_run_underline(target_run, font.underline, source_run=source_run)
                except Exception:
                    pass

            def segment_placeholder_style(positions, start, end):
                styles = {
                    'underline': False,
                    'bold': False,
                    'italic': False,
                    'ascii_name': None,
                    'east_asia_name': None,
                    'h_ansi_name': None,
                    'cs_name': None,
                    'font_size': None,
                    'color': None,
                    'highlight': None,
                }
                for s, e, run in positions:
                    if e > start and s < end:
                        try:
                            r_fonts = get_run_r_fonts(run)
                            if r_fonts is not None:
                                styles['ascii_name'] = r_fonts.get(qn('w:ascii'))
                                styles['east_asia_name'] = r_fonts.get(qn('w:eastAsia'))
                                styles['h_ansi_name'] = r_fonts.get(qn('w:hAnsi'))
                                styles['cs_name'] = r_fonts.get(qn('w:cs'))
                            run_font_name = run.font.name
                            if run_font_name is not None:
                                styles['ascii_name'] = run_font_name
                                if styles['east_asia_name'] is None:
                                    styles['east_asia_name'] = run_font_name
                            if run.font.size is not None:
                                styles['font_size'] = run.font.size
                            if run.font.color and run.font.color.rgb is not None:
                                styles['color'] = run.font.color.rgb
                            if run.underline:
                                styles['underline'] = run.underline
                            if run.bold:
                                styles['bold'] = True
                            if run.italic:
                                styles['italic'] = True
                            if run.font.highlight_color is not None:
                                styles['highlight'] = run.font.highlight_color
                        except Exception:
                            continue
                return styles

            def replace_paragraph_runs(paragraph):
                runs = list(paragraph.runs)
                if not runs:
                    return

                full_text = ''.join(run.text for run in runs)
                replaced_text = replace_placeholders(full_text)
                if replaced_text == full_text:
                    return

                matches = list(placeholder_pattern.finditer(full_text))
                if not matches:
                    return

                positions = get_run_segments(paragraph)
                for run in runs:
                    run._element.getparent().remove(run._element)

                last_index = 0
                for match in matches:
                    if match.start() > last_index:
                        text_segment = full_text[last_index:match.start()]
                        seg_pos = next((p for p in positions if p[0] <= last_index and p[1] > last_index), None)
                        if seg_pos:
                            segment_run = paragraph.add_run(replace_placeholders(text_segment))
                            copy_run_style(seg_pos[2], segment_run)
                        else:
                            segment_run = paragraph.add_run(replace_placeholders(text_segment))
                            copy_run_style(runs[0], segment_run)

                    replacement_text = replace_placeholders(match.group())
                    placeholder_style = segment_placeholder_style(positions, match.start(), match.end())
                    placeholder_run = paragraph.add_run(replacement_text)

                    if placeholder_style['ascii_name'] is not None:
                        try:
                            set_run_font(
                                placeholder_run,
                                ascii_name=placeholder_style['ascii_name'],
                                east_asia_name=placeholder_style['east_asia_name'],
                                h_ansi_name=placeholder_style['h_ansi_name'],
                                cs_name=placeholder_style['cs_name'],
                            )
                        except Exception:
                            pass
                    if placeholder_style['font_size'] is not None:
                        try:
                            placeholder_run.font.size = placeholder_style['font_size']
                        except Exception:
                            pass
                    if placeholder_style['color'] is not None:
                        try:
                            placeholder_run.font.color.rgb = placeholder_style['color']
                        except Exception:
                            pass
                    try:
                        set_run_underline(placeholder_run, placeholder_style['underline'])
                    except Exception:
                        pass
                    if placeholder_style['bold']:
                        try:
                            placeholder_run.bold = True
                        except Exception:
                            pass
                    if placeholder_style['italic']:
                        try:
                            placeholder_run.italic = True
                        except Exception:
                            pass
                    if placeholder_style['highlight'] is not None:
                        try:
                            placeholder_run.font.highlight_color = placeholder_style['highlight']
                        except Exception:
                            pass

                    last_index = match.end()

                if last_index < len(full_text):
                    for s, e, run in positions:
                        if e > last_index:
                            local_start = max(s, last_index) - s
                            run_text = run.text[local_start:]
                            if run_text:
                                tail_run = paragraph.add_run(replace_placeholders(run_text))
                                copy_run_style(run, tail_run)

            for paragraph in doc.paragraphs:
                replace_paragraph_runs(paragraph)
                
                # 设置段落格式以防止自动断行
                p_pr = paragraph._element.get_or_add_pPr()
                
                # 添加keepNext（与下一段保持在一起）
                keep_next = p_pr.find(W_NS + 'keepNext')
                if keep_next is None:
                    from lxml import etree
                    keep_next = etree.SubElement(p_pr, W_NS + 'keepNext')
                keep_next.set(qn('w:val'), '0')
                
                # 添加keepLines（段落内保持连续，不分页）
                keep_lines = p_pr.find(W_NS + 'keepLines')
                if keep_lines is None:
                    from lxml import etree
                    keep_lines = etree.SubElement(p_pr, W_NS + 'keepLines')
                keep_lines.set(qn('w:val'), '0')
                
                # 添加 widowControl（寡妇/孤儿控制）
                widow_control = p_pr.find(W_NS + 'widowControl')
                if widow_control is None:
                    from lxml import etree
                    widow_control = etree.SubElement(p_pr, W_NS + 'widowControl')
                widow_control.set(qn('w:val'), '1')
                
                # 设置段落不允许断字（no hyphenation）
                # 通过设置字间距来避免自动断行
                for run in paragraph.runs:
                    r_pr = run._element.get_or_add_rPr()
                    # 设置禁止断字
                    no_hyphen = r_pr.find(W_NS + 'noHyphen')
                    if no_hyphen is None:
                        from lxml import etree
                        no_hyphen = etree.SubElement(r_pr, W_NS + 'noHyphen')
                    no_hyphen.set(qn('w:val'), '1')

            # 替换表格中的文本，按段落处理以支持跨run占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_paragraph_runs(paragraph)
                            
                            # 同样为表格中的段落设置格式
                            p_pr = paragraph._element.get_or_add_pPr()
                            
                            keep_next = p_pr.find(W_NS + 'keepNext')
                            if keep_next is None:
                                from lxml import etree
                                keep_next = etree.SubElement(p_pr, W_NS + 'keepNext')
                            keep_next.set(qn('w:val'), '0')
                            
                            keep_lines = p_pr.find(W_NS + 'keepLines')
                            if keep_lines is None:
                                from lxml import etree
                                keep_lines = etree.SubElement(p_pr, W_NS + 'keepLines')
                            keep_lines.set(qn('w:val'), '0')
                            
                            widow_control = p_pr.find(W_NS + 'widowControl')
                            if widow_control is None:
                                from lxml import etree
                                widow_control = etree.SubElement(p_pr, W_NS + 'widowControl')
                            widow_control.set(qn('w:val'), '1')
                            
                            for run in paragraph.runs:
                                r_pr = run._element.get_or_add_rPr()
                                no_hyphen = r_pr.find(W_NS + 'noHyphen')
                                if no_hyphen is None:
                                    from lxml import etree
                                    no_hyphen = etree.SubElement(r_pr, W_NS + 'noHyphen')
                                no_hyphen.set(qn('w:val'), '1')

            # 保存DOCX文件
            output_path = os.path.join(self.output_dir, f"{output_name}.docx")
            os.makedirs(self.output_dir, exist_ok=True)
            doc.save(output_path)

            print(f"生成DOCX文件: {output_path}")
            return output_path

        except Exception as e:
            print(f"填充DOCX模板失败: {e}")
            raise

    def convert_to_png(self, docx_path: str, png_output_dir: str, dpi: int = 300) -> list:
        """将DOCX转换为高清PNG图片

        Args:
            docx_path: DOCX文件路径
            png_output_dir: PNG输出目录
            dpi: 分辨率（dpi）

        Returns:
            生成的PNG文件路径列表
        """
        try:
            os.makedirs(png_output_dir, exist_ok=True)

            # 获取文件名（不含扩展名）
            base_name = Path(docx_path).stem

            # 使用LibreOffice将DOCX转换为PDF（中间步骤）
            pdf_path = os.path.join(png_output_dir, f"{base_name}.pdf")

            pdf_created = False

            # 使用LibreOffice转换为PDF
            soffice_exe = find_soffice()
            if not soffice_exe:
                print("未找到 LibreOffice，跳过 PNG 转换。（安装 LibreOffice 或设置 LO_PATH 环境变量）")
                return []

            # 创建LibreOffice用户配置目录（headless模式需要）
            user_profile = os.path.join(png_output_dir, ".libreoffice_profile")
            os.makedirs(user_profile, exist_ok=True)
            
            # 创建LibreOffice配置文件以优化排版，避免自动断行
            config_dir = os.path.join(user_profile, "user", "config")
            os.makedirs(config_dir, exist_ok=True)
            
            # 创建registrymodifications.xcu配置文件
            config_content = '''<?xml version="1.0" encoding="UTF-8"?>
<oor:items xmlns:oor="http://openoffice.org/2001/registry" xmlns:xs="http://www.w3.org/2001/XMLSchema" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">
  <item oor:path="/org.openoffice.Office.Writer/Content/Visibility">
    <prop oor:name="ShowParagraphEnd" oor:op="fuse"><value>false</value></prop>
  </item>
  <item oor:path="/org.openoffice.Office.Writer/Content/Formatting">
    <prop oor:name="ShowHyphens" oor:op="fuse"><value>false</value></prop>
  </item>
  <item oor:path="/org.openoffice.Office.Writer/Compatibility">
    <prop oor:name="UseFormerLineSpacing" oor:op="fuse"><value>true</value></prop>
    <prop oor:name="AddParaSpacingToTableCells" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="UseFormerObjectPos" oor:op="fuse"><value>true</value></prop>
    <prop oor:name="UseFormerTextWrapping" oor:op="fuse"><value>true</value></prop>
    <prop oor:name="ConsiderWrapOnObjPos" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="AddSpacingBetweenAsianText" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="AddSpacingBetweenAsianAndNumber" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="AddSpacingBetweenAsianAndLatin" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="SingleLineSpacingAtBottom" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="TabStopAtPageStart" oor:op="fuse"><value>true</value></prop>
  </item>
  <item oor:path="/org.openoffice.Office.Writer/Print">
    <prop oor:name="PrintTextPlaceholder" oor:op="fuse"><value>false</value></prop>
  </item>
  <item oor:path="/org.openoffice.Office.Writer/Layout">
    <prop oor:name="TextBoundaries" oor:op="fuse"><value>false</value></prop>
    <prop oor:name="TextFields" oor:op="fuse"><value>false</value></prop>
  </item>
</oor:items>
'''
            config_file = os.path.join(config_dir, "registrymodifications.xcu")
            with open(config_file, 'w', encoding='utf-8') as f:
                f.write(config_content)

            # 尝试多种转换方式
            convert_success = False
            
            # 方式1: 使用标准pdf filter，添加优化参数
            try:
                cmd = [
                    soffice_exe,
                    "--headless",
                    "--norestore",
                    "--invisible",
                    "--nodefault",
                    "--nolockcheck",
                    "--convert-to", "pdf:writer_pdf_Export",
                    "--outdir", png_output_dir,
                    docx_path
                ]
                env = os.environ.copy()
                env["HOME"] = user_profile
                
                # 设置环境变量以优化排版
                env["SAL_VCL_FORCEX11CAIRO"] = "0"
                env["SAL_USE_VCLPLUGIN"] = "gen"  # 使用通用后端，避免X11渲染问题
                env["DISPLAY"] = ""  # 确保不使用X11
                env["QT_QPA_PLATFORM"] = "offscreen"  # 禁用Qt平台
                
                subprocess.run(cmd, check=True, capture_output=True, timeout=60, env=env)
                
                if os.path.exists(pdf_path):
                    print(f"已转换为PDF（LibreOffice）: {pdf_path}")
                    convert_success = True
            except Exception as e:
                print(f"LibreOffice转换失败(方式1): {e}")

            # 方式2: 使用默认pdf filter
            if not convert_success:
                try:
                    cmd = [
                        soffice_exe,
                        "--headless",
                        "--norestore",
                        "--invisible",
                        "--convert-to", "pdf",
                        "--outdir", png_output_dir,
                        docx_path
                    ]
                    env = os.environ.copy()
                    env["HOME"] = user_profile
                    subprocess.run(cmd, check=True, capture_output=True, timeout=60, env=env)
                    
                    if os.path.exists(pdf_path):
                        print(f"已转换为PDF（LibreOffice）: {pdf_path}")
                        convert_success = True
                except Exception as e:
                    print(f"LibreOffice转换失败(方式2): {e}")

            pdf_created = convert_success

            # 使用pdf2image将PDF转换为PNG
            if pdf_created and os.path.exists(pdf_path):
                from pdf2image import convert_from_path
                from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError

                # 在 Linux/Docker 环境中显式设置 pdftoppm 路径
                poppler_path = None
                if sys.platform != 'win32':
                    for candidate in ['/usr/bin/pdftoppm', '/usr/local/bin/pdftoppm']:
                        if os.path.isfile(candidate):
                            poppler_path = os.path.dirname(candidate)
                            break
                    if not poppler_path:
                        import subprocess as _sp
                        try:
                            result = _sp.run(['which', 'pdftoppm'], capture_output=True, text=True)
                            if result.returncode == 0 and result.stdout.strip():
                                poppler_path = os.path.dirname(result.stdout.strip())
                        except Exception:
                            pass

                try:
                    images = convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path)
                except PDFInfoNotInstalledError:
                    print("poppler-utils 未安装，跳过 PNG 转换")
                    return []
                except Exception as e:
                    print(f"PDF转PNG失败: {e}")
                    return []

                png_paths = []

                for i, image in enumerate(images, 1):
                    png_path = os.path.join(png_output_dir, f"{base_name}_page{i}.png")
                    image.save(png_path, 'PNG')
                    png_paths.append(png_path)
                    print(f"生成PNG文件: {png_path}")

                return png_paths
            else:
                print(f"PDF文件不存在，无法转换为PNG")
                return []

        except Exception as e:
            print(f"转换为PNG失败: {e}")
            # 返回空列表而不抛出异常，允许程序继续
            return []


if __name__ == "__main__":
    handler = DOCXHandler("gov_invitation_template.docx", "output_documents")
    data = {
        "姓名": "李四",
        "单位": "测试机构",
        "职务": "主任"
    }
    docx_file = handler.fill_template(data, "邀请函_李四")
    # handler.convert_to_png(docx_file, "output_images")
